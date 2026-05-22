# -*- coding: utf-8 -*-
"""
골든듀 통합 정책 안내서 생성 스크립트 (v2 - 인수인계용 상세본)
출력: C:\\Users\\milvus-0\\Downloads\\골든듀_통합정책본.docx (덮어쓰기)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===== 스타일 헬퍼 =====
KOREAN_FONT = "맑은 고딕"
NAVY = RGBColor(0x1F, 0x3A, 0x6B)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
GRAY_BG = "F2F2F2"
BLUE_BG = "DCE6F1"
RED_BG = "FBE4E4"
GREEN_BG = "E2EFDA"
YELLOW_BG = "FFF2CC"


def set_run_font(run, name=KOREAN_FONT, size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="BFBFBF", size="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    if level == 0:
        set_run_font(run, size=22, bold=True, color=NAVY)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        set_run_font(run, size=16, bold=True, color=NAVY)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=NAVY)
    else:
        set_run_font(run, size=11, bold=True)
    return p


def add_para(doc, text, size=10, bold=False, indent=0, color=None, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, size=10, indent=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_numbered(doc, text, size=10, indent=0.5):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_callout(doc, title, body, kind="info"):
    fill = {"info": BLUE_BG, "warn": RED_BG, "ok": GREEN_BG, "note": YELLOW_BG}[kind]
    title_color = {"info": NAVY, "warn": RED, "ok": GREEN, "note": NAVY}[kind]

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, fill)
    set_cell_borders(cell, color="BFBFBF")
    cell.width = Cm(16)

    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    r = p_title.add_run(title)
    set_run_font(r, size=10, bold=True, color=title_color)

    for line in body.split("\n"):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        rr = p.add_run(line)
        set_run_font(rr, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, "1F3A6B")
        set_cell_borders(cell, color="1F3A6B")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for ri, row in enumerate(rows):
        is_alt = ri % 2 == 1
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if is_alt:
                set_cell_bg(cell, GRAY_BG)
            set_cell_borders(cell, color="BFBFBF")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            r = p.add_run(val)
            set_run_font(r, size=10)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_break(doc):
    doc.add_page_break()


def add_header_footer(doc, title):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(title)
    set_run_font(r, size=9, color=NAVY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("골든듀 통합 정책 안내서  |  발행일 2026-05-21  |  대외비")
    set_run_font(fr, size=8, color=RGBColor(0x80, 0x80, 0x80))


# ===========================================================
#                       문서 생성 시작
# ===========================================================
doc = Document()

style = doc.styles["Normal"]
style.font.name = KOREAN_FONT
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)

add_header_footer(doc, "골든듀 통합 정책 안내서")

# ===== 표지 =====
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(120)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("업무 기준 안내 문서")
set_run_font(r, size=11, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
r = p.add_run("골든듀 통합 정책 안내서")
set_run_font(r, size=26, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Goldendew Integrated Policy Guide")
set_run_font(r, size=12, color=RGBColor(0x60, 0x60, 0x60))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("멤버십 · 포인트 · 프로모션 · 쿠폰")
set_run_font(r, size=12, color=NAVY)

add_callout(
    doc,
    "개요",
    "본 문서는 골든듀 CRM 시스템에서 운영되는 멤버십, 포인트, 프로모션, 쿠폰 정책을 통합하여 정리한 인수인계용 업무 기준 안내서입니다.\n"
    "각 정책은 현재 운영 중인 Salesforce/Apex 구현 로직과 사용자 매뉴얼을 기준으로 도출되었으며, "
    "실무자가 정책 의사결정과 고객 응대 시 참고할 수 있도록 작성되었습니다.\n"
    "정책 적용 기준일은 본 문서 발행일이며, 시스템 변경 시 별도 개정본을 발행합니다.",
)

add_page_break(doc)

# ===== 목차 =====
add_heading(doc, "목차", level=1)
toc = [
    "1. 멤버십(회원 등급) 정책",
    "   1-1. 등급 체계 (GIP / VIP / VVIP / SVIP)",
    "   1-2. 등급 산정 기준 — 3년 누적 구매금액",
    "   1-3. 3년 누적 금액 계산 규칙 (포함/제외)",
    "   1-4. 등급 전환 스케줄 및 배치 파이프라인",
    "   1-5. 등급별 혜택 — 적립률, 사은 보너스, 등급 쿠폰",
    "   1-6. 예외 및 엣지 케이스",
    "2. 포인트 정책",
    "   2-1. 포인트 종류 — 7개 SubType 전체 분류",
    "   2-2. 등급별 구매 시 적립 비율",
    "   2-3. 적립(Credit) 규칙 — 구매/수동/API/프로모션",
    "   2-4. 차감(Debit·사용) 규칙 및 우선순위",
    "   2-5. 환불(Return) 규칙",
    "   2-6. 유효기간 및 소멸 정책",
    "   2-7. 마이그레이션 이전 판매 건 처리 (기타포인트)",
    "   2-8. 원장(Ledger) 구조 및 집계",
    "   2-9. ERP 연동 흐름",
    "3. 프로모션 정책",
    "   3-1. 프로모션 구성 요소",
    "   3-2. 적용 조건 (기간·매장·상품·할인율)",
    "   3-3. 쿠폰 사용 여부 체크 항목",
    "   3-4. 프로모션 포인트 적립·차감 규칙",
    "   3-5. 적립 불가 포인트 유형 설정",
    "   3-6. 트리거 및 자동화",
    "   3-7. ERP 연동 및 업로드",
    "4. 쿠폰 정책",
    "   4-1. 쿠폰 적용 도움창 기본 조회 조건",
    "   4-2. 프로모션 미적용 거래",
    "   4-3. 프로모션 적용 거래",
    "   4-4. 쿠폰 사용 조건별 확인 기준",
    "   4-5. 쿠폰이 조회되지 않을 때 확인 순서",
    "   4-6. 대표 예시",
    "5. 한 줄 요약",
]
for line in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(line)
    set_run_font(r, size=11)

add_page_break(doc)

# ============================================================
# 1. 멤버십 정책
# ============================================================
add_heading(doc, "1. 멤버십(회원 등급) 정책", level=1)
add_para(
    doc,
    "골든듀 멤버십은 고객의 3년 누적 구매 실적에 따라 등급을 차등 부여하고, 등급별 차등 혜택(포인트 적립률, 사은 보너스, 등급 쿠폰 등)을 제공하는 로열티 프로그램입니다. "
    "등급 산정은 자동화된 배치를 통해 매년 정해진 기준일에 일괄 전환되며, ERP와 양방향으로 동기화됩니다.",
)

add_heading(doc, "1-1. 등급 체계 (GIP / VIP / VVIP / SVIP)", level=2)
add_para(doc, "회원 등급은 Contact 객체의 Grade__c 필드(Picklist)로 관리되며, LoyaltyTier 객체와 코드값으로 매핑됩니다. 등급은 4단계로 구성되어 있습니다.")
add_table(
    doc,
    headers=["등급 순서", "등급명", "기본 진입 기준", "비고"],
    rows=[
        ["1", "GIP", "신규 가입 시 기본 부여", "전체 회원의 대부분이 속한 기본 등급"],
        ["2", "VIP", "3년 누적 1천만 원 이상", "임계값 기반 자동 승급"],
        ["3", "VVIP", "VIP 상위 등급", "LoyaltyTier.MinimumEligibleBalance에 따른 임계값 기반"],
        ["4", "SVIP", "최상위 등급", "VVIP 상위, 임계값은 LoyaltyTier 설정값 참조"],
    ],
    col_widths=[2.5, 2.5, 5, 6],
)
add_bullet(doc, "현재 등급(Grade__c): 회원이 현재 시점에 보유한 등급")
add_bullet(doc, "예상 등급(ExpectedGrade__c): 다음 전환일에 적용될 예정 등급")
add_bullet(doc, "이전 등급(PreviousGrade__c): 등급 변경 추적용 직전 등급")
add_callout(
    doc,
    "참고",
    "VVIP·SVIP의 정확한 임계 금액은 LoyaltyTier.MinimumEligibleBalance(등급 최소 임계금액) 설정값에 의해 결정됩니다. "
    "임계값 변경은 LoyaltyTier 레코드 직접 수정으로 즉시 반영되지만, 실제 회원 등급 갱신은 다음 전환일(2월 1일)에 일괄 적용됩니다.",
    kind="note",
)

add_heading(doc, "1-2. 등급 산정 기준 — 3년 누적 구매금액", level=2)
add_para(doc, "등급 산정의 핵심 지표는 Contact.TierQualifyingSpend3Y__c(3년 누적 순판매액)이며, 다음 규칙으로 운영됩니다.")
add_bullet(doc, "산식: 순판매액 = 판매액 − 반품액 (단가 기준 NetSalesUnitPrice__c)")
add_bullet(doc, "비교 대상: 각 LoyaltyTier.MinimumEligibleBalance(등급 최소 임계금액)")
add_bullet(doc, "매칭 방식: 임계값 DESC 정렬 후 첫 매칭 (최고 등급 우선 배정)")
add_bullet(doc, "TierQualifyingSpend3Y__c = NULL → 0으로 간주 → 최저 등급(GIP) 배정")

add_heading(doc, "1-3. 3년 누적 금액 계산 규칙 (포함/제외)", level=2)
add_callout(
    doc,
    "전환 기준일",
    "매년 2월 1일 (코드 상수: CONVERSION_MONTH=2, CONVERSION_DAY=1)\n"
    "예시: 2026년 기준이라면 2024-01-01 ~ 2026-01-31 기간의 순판매액 합산\n"
    "→ 즉, '직전 3개년'의 1월 1일부터 전환일 직전(1월 31일)까지를 합산합니다.",
    kind="info",
)
add_para(doc, "■ 포함 대상", bold=True, after=2)
add_bullet(doc, "Account.StoreType__c IN ('11', '40') 또는 StoreTypeFour__c='01' 매장의 판매 성공(saleSuccess) / 반품 성공(returnSuccess) 주문")
add_bullet(doc, "NetSalesUnitPrice__c 기준 순판매액 (판매액 − 반품액)")
add_para(doc, "")
add_para(doc, "■ 제외 대상", bold=True, after=2)
add_bullet(doc, "골드바 (ProductCategory='골드바')")
add_bullet(doc, "수선 / 각인 / 위약금 / 재감비 / 재인쇄비 상품")
add_para(doc, "")

add_heading(doc, "1-4. 등급 전환 스케줄 및 배치 파이프라인", level=2)
add_para(doc, "등급 전환은 다음 3단계 파이프라인으로 자동 실행됩니다.")
add_numbered(doc, "TierQualifyingSpend3YRecalcBatch — 3년 누적 순판매액 재계산 → Contact.TierQualifyingSpend3Y__c 갱신")
add_numbered(doc, "ExpectedGradeCalculationBatch — 임계값 비교(DESC 정렬 최고 등급 매칭) → Contact.ExpectedGrade__c 산출")
add_numbered(doc, "LoyaltyMemberTierUpdateBatch — Grade__c ← ExpectedGrade__c 승격, ERP 동기화 enqueue")
add_para(doc, "")
add_callout(
    doc,
    "스케줄러 정보",
    "LoyaltyMemberTierUpdateScheduler: 매년 2월 1일 실행, 100명 단위 청크 처리\n"
    "대상 조건: ExpectedGrade__c IS NOT NULL AND MemberDropDate__c IS NULL AND Grade__c IS NOT NULL\n"
    "ERP 연동: 등급 변경 성공 시 ErpMemberJob(UPDATE_GRADE) enqueue → MemberNo__c + Grade__c 전송",
    kind="info",
)

add_heading(doc, "1-5. 등급별 혜택 — 적립률, 사은 보너스, 등급 쿠폰", level=2)
add_para(doc, "■ 등급별 구매 포인트 적립 비율 (사용자 매뉴얼 기준)", bold=True, after=2)
add_table(
    doc,
    headers=["등급", "구매포인트 적립률", "사은포인트 적립률", "비고"],
    rows=[
        ["GIP", "구매 금액의 1%", "구매 금액의 1%", "기본 등급"],
        ["VIP", "구매 금액의 1.5%", "구매 금액의 2%", "1천만 원 이상 누적 시 승급"],
        ["VVIP", "Benefit 설정값 참조", "Benefit 설정값 참조", "LoyaltyTierBenefit/Benefit 객체에 등록된 보상 정의에 따름"],
        ["SVIP", "Benefit 설정값 참조", "Benefit 설정값 참조", "동일 — 최상위 등급 보상 정의에 따름"],
    ],
    col_widths=[2.5, 4, 4, 5.5],
)
add_callout(
    doc,
    "적립률 운영 메모",
    "구매포인트와 사은포인트는 별도 SubType으로 동시 적립됩니다 (즉, 구매 시 두 종류의 포인트가 함께 쌓임).\n"
    "VVIP / SVIP의 정확한 적립률은 사용자 매뉴얼에 명시되어 있지 않으며, '설정 > 멤버십'에서 해당 등급의 '회원 혜택 → 보상 이력' 탭을 확인하면 됩니다. "
    "보상은 LoyaltyTierBenefit / Benefit 객체의 '보상 정의 추가' 화면에서 등록·수정합니다.",
    kind="note",
)

add_para(doc, "■ 등급 쿠폰 발급 (IssueTierCouponBatch)", bold=True, after=2)
add_numbered(doc, "대상 LoyaltyTier와 VoucherDefinition, IssueDate를 파라미터로 배치 실행")
add_numbered(doc, "TransactionJournal 사전 생성 (Type=Accrual, SubType=Manual Coupon Issuance)")
add_numbered(doc, "LoyaltyVoucherService.issueVoucherV2() 호출하여 쿠폰 일괄 지급")
add_numbered(doc, "IsUpgrade__c=TRUE 쿠폰은 FaceValue=0으로 정규화 (등급 승급 쿠폰)")
add_para(doc, "")
add_callout(
    doc,
    "쿠폰 등급 제한 (VoucherLimit__c)",
    "VoucherLimit__c(Type__c='GradeIncludeList')로 쿠폰별 사용 가능 등급을 제한합니다.\n"
    "VoucherDefinitionId__c + Value__c(등급 코드) 매핑을 검증하여 허용된 등급의 회원에게만 쿠폰이 노출됩니다.",
    kind="info",
)

add_heading(doc, "1-6. 예외 및 엣지 케이스", level=2)
add_bullet(doc, "탈퇴 회원: MemberDropDate__c IS NOT NULL → 모든 등급 배치에서 제외")
add_bullet(doc, "LoyaltyProgramMember 미존재 회원: 등급 전환 실패 처리")
add_bullet(doc, "TierQualifyingSpend3Y__c = NULL → 0으로 간주 → 최저 등급(GIP) 배정")
add_bullet(doc, "쿠폰 발급 시 ExpirationDate < Date.today() 인 쿠폰은 지급 불가")
add_bullet(doc, "배치는 Database.update(..., false)로 부분 성공 허용 (SUCCESS / PARTIAL / ERROR 분류)")
add_para(doc, "")
add_callout(
    doc,
    "다운그레이드 규칙",
    "현재 시스템은 Grade__c ← ExpectedGrade__c 일방향 동기 구조로, 3년 누적 금액 감소 시 등급 재산정은 매년 2월 1일 전환일에만 발생합니다. "
    "연중 임의 시점의 다운그레이드 처리 규칙은 코드상 별도 정의되어 있지 않습니다.",
    kind="warn",
)

add_page_break(doc)

# ============================================================
# 2. 포인트 정책
# ============================================================
add_heading(doc, "2. 포인트 정책", level=1)
add_para(
    doc,
    "골든듀 포인트는 'Golden Dew' 로열티 프로그램의 '포인트' 통화로 운영되며, 적립·차감·소멸 내역이 LoyaltyLedger에 모두 기록됩니다. "
    "포인트는 총 7개의 SubType(종류)으로 구분되며, 종류별로 적립 사유·사용 우선순위·만료 정책이 다르게 적용되고 ERP와 실시간 동기화됩니다.",
)

add_heading(doc, "2-1. 포인트 종류 — 7개 SubType 전체 분류", level=2)
add_para(doc, "포인트는 LoyaltyPgmCrcySubtype.PointTypeCode__c 코드(2자리)로 구분되며, 사용자 매뉴얼의 '포인트 지급' 화면에서 선택 가능한 7개 종류와 1:1 대응됩니다.")
add_table(
    doc,
    headers=["코드", "구분(SubType명)", "적립 사유", "비고"],
    rows=[
        ["01", "구매포인트", "거래 등급 비율에 따른 자동 적립", "정상 구매 거래의 기본 적립"],
        ["02", "기념일포인트", "생일·결혼기념일 등 기념일 보너스", "수동·자동 지급 모두 가능"],
        ["03", "승급포인트", "등급 승급 시 보너스 적립", "등급 전환 배치와 연동"],
        ["04", "기타포인트", "마이그레이션 오픈 일자 이전 판매 건에 대한 포인트", "구 시스템 이관 데이터 전용 (아래 2-7 참조)"],
        ["05", "에코포인트", "API 강제 지급(적립) 처리 전용", "외부 API 요청 기반의 강제 적립/차감/소멸"],
        ["06", "사은포인트", "등급 보상으로 구매 시 사은 보너스 적립", "구매포인트와 동시에 별도로 적립"],
        ["07", "이벤트포인트", "이벤트성 수동 지급 포인트", "수동 지급 시 가장 일반적으로 사용"],
    ],
    col_widths=[1.5, 3, 6, 5.5],
)
add_callout(
    doc,
    "핵심 구분 — 에코포인트 vs 이벤트포인트",
    "두 포인트는 자주 혼동되지만 명확히 다른 SubType입니다.\n"
    "  · 에코포인트(05): 외부 시스템에서 API로 강제 지급/차감/소멸하는 운영성 포인트. "
    "관리자 화면의 '포인트 지급' 수동 입력으로는 사용하지 않습니다.\n"
    "  · 이벤트포인트(07): 마케팅·CS·이벤트 운영 등 관리자가 수동으로 직접 지급할 때 사용하는 일반 보너스 포인트.\n"
    "ERP 별첨 매핑상 'API 요청에 의한 포인트 강제 지급(적립) 처리 = 에코포인트 적립', "
    "'포인트 수동 지급 = 이벤트 포인트 지급'으로 정의되어 있습니다.",
    kind="warn",
)

add_heading(doc, "2-2. 등급별 구매 시 적립 비율", level=2)
add_para(
    doc,
    "구매 거래에서 자동 적립되는 포인트는 '구매포인트(01)'와 '사은포인트(06)' 두 가지이며, "
    "등급별로 적립률이 다르게 적용됩니다. 비율은 LoyaltyTierBenefit / Benefit 객체의 '보상 정의'에 등록되어 운영됩니다.",
)
add_table(
    doc,
    headers=["등급", "구매포인트(01)", "사은포인트(06)", "구매 100만 원 시 적립 예시"],
    rows=[
        ["GIP", "1.0%", "1.0%", "구매 10,000P + 사은 10,000P"],
        ["VIP", "1.5%", "2.0%", "구매 15,000P + 사은 20,000P"],
        ["VVIP", "Benefit 설정값", "Benefit 설정값", "관리자 화면 '설정 > 멤버십 > VVIP > 보상 이력'에서 확인"],
        ["SVIP", "Benefit 설정값", "Benefit 설정값", "관리자 화면 '설정 > 멤버십 > SVIP > 보상 이력'에서 확인"],
    ],
    col_widths=[2.5, 3.5, 3.5, 6.5],
)
add_callout(
    doc,
    "운영 메모",
    "보상 비율은 LoyaltyTierBenefit / Benefit 객체에서 관리되며, 변경 시 즉시 신규 거래에 반영됩니다.\n"
    "VVIP / SVIP의 비율은 코드 상수가 아닌 데이터 설정값이므로, 본 문서 작성 시점의 비율을 별도 명기하지 않습니다. "
    "정확한 운영 비율은 '설정 - 멤버십' 메뉴에서 각 등급의 '회원 혜택' 카드를 확인하세요.",
    kind="note",
)

add_heading(doc, "2-3. 적립(Credit) 규칙 — 구매/수동/API/프로모션", level=2)
add_para(doc, "■ 구매 시 자동 적립", bold=True, after=2)
add_bullet(doc, "기본: GdPointServiceImpl.creditPointsOnPurchase() — 거래 등급 비율 기반 자동 적립 (구매포인트 + 사은포인트 동시)")
add_bullet(doc, "V2: creditPointsOnPurchaseV2() — excludeAmount 파라미터로 적립 제외 금액 지정 가능 (프로모션 등 일부 금액 제외)")
add_bullet(doc, "프로모션 포인트: creditPromotionPointsOnPurchase() — 프로모션 별도 적립")
add_para(doc, "")
add_para(doc, "■ 적립 대상 판단", bold=True, after=2)
add_bullet(doc, "isAccrualEligible(TransactionJournal tj) 메서드로 적립 가능 여부 판단")
add_bullet(doc, "아울렛 매장(StoreType: 20, 25, 33)은 적립 제외")
add_para(doc, "")
add_para(doc, "■ 수동 적립 (관리자 화면 '포인트 지급' 기능)", bold=True, after=2)
add_bullet(doc, "회원 360 > 회원 요약 > '포인트 지급' 버튼 → 개인 단위 지급")
add_bullet(doc, "설정 - 포인트 > 해당 포인트 유형 > '지급 대상 업로드' → CSV(NO_MBR 컬럼)로 그룹 일괄 지급 (최대 1만 건 권장)")
add_bullet(doc, "그룹 지급 시 만료 일자: 기본 2년, 수동 지정 가능")
add_bullet(doc, "지급 가능한 포인트 종류: 기념일포인트, 승급포인트, 기타포인트, 에코포인트, 사은포인트, 구매포인트, 이벤트포인트 (총 7종)")
add_para(doc, "")
add_para(doc, "■ API 강제 적립 (에코포인트)", bold=True, after=2)
add_bullet(doc, "외부 시스템에서 호출하는 API 기반 강제 적립 — 모두 에코포인트(05)로 처리")
add_bullet(doc, "관련 상수: JST_API_FORCED_CREDIT = 'Api Forced Credit'")
add_bullet(doc, "ERP 연동 시 loyaltyProgramCurrencySubtypeName='에코포인트', additionalNotes='이벤트포인트'로 전송 (코드 라인 GdPointServiceImpl 1781, 1885)")
add_para(doc, "")
add_callout(
    doc,
    "적립 만료일",
    "적립 시점에 EscrowPointsCreditDate 필드를 통해 만료일을 함께 기록합니다. "
    "이후 LoyaltyPointExpirationBatch가 만료일 기준으로 자동 소멸 처리합니다. "
    "그룹 지급 화면에서도 만료일은 기본 2년이며 수동 변경 가능합니다.",
    kind="info",
)

add_heading(doc, "2-4. 차감(Debit·사용) 규칙 및 우선순위", level=2)
add_para(doc, "■ 차감 메서드", bold=True, after=2)
add_bullet(doc, "표준 차감: GdPointServiceImpl.debitPoint(TransactionJournal tj)")
add_bullet(doc, "판매 수정 전용: debitPointForSaleUpdate() — 같은 판매에서 적립된 포인트 제외, 커스텀 DML 방식")
add_bullet(doc, "특정 차감: debitSpecificPoints / V2 / V3 — LoyaltyLedger ID 기반 지정 차감")
add_para(doc, "")
add_para(doc, "■ 차감 우선순위 — 만료 임박 → SubType 우선순위 순", bold=True, after=2)
add_numbered(doc, "1순위: 만료일자 오름차순 (ExpiryDate ASC) — 먼저 만료될 포인트부터 사용")
add_numbered(doc, "2순위: 포인트 SubType의 UsagePriorityNumber (값이 클수록 우선)")
add_para(doc, "")
add_table(
    doc,
    headers=["SubType", "UsagePriorityNumber", "차감 우선순위(예시 정렬 결과)"],
    rows=[
        ["구매포인트", "7", "1순위 (가장 먼저 차감)"],
        ["사은포인트", "5", "2순위"],
        ["이벤트포인트", "1", "3순위 (가장 나중 차감)"],
    ],
    col_widths=[5, 5, 6],
)
add_callout(
    doc,
    "잔액 부족 시 에러 메시지",
    '"포인트 잔액이 부족합니다. 사용 가능 금액: {available}P (요청: {requested}P)"',
    kind="warn",
)

add_heading(doc, "2-5. 환불(Return) 규칙", level=2)
add_bullet(doc, "에코포인트 복원: restoreEcoPointsFromReturn(Id orderId) — 전액 반품 시 에코포인트 재적립")
add_bullet(doc, "적립 취소: ECO_POINT_CANCEL_NOTE = 'FULL_RETURN_ECO_POINT_CANCEL'")
add_bullet(doc, "포인트 사용 복구: 반품 입금 삭제 시 RETURN_DEPOSIT_DELETE_POINT_RESTORE_NOTE 처리")
add_bullet(doc, "구매 시 적립된 구매포인트·사은포인트는 반품 시 적립 취소(복원 포인트)로 처리")
add_para(doc, "")
add_callout(
    doc,
    "부분 반품 처리",
    "현재 부분 반품 API(GdReturnPointApiControllerV1)는 비활성화 상태입니다. "
    "운영 정책 확정 후 활성화 예정이며, 그 전까지는 전액 반품 케이스만 자동 처리됩니다.",
    kind="warn",
)

add_heading(doc, "2-6. 유효기간 및 소멸 정책", level=2)
add_callout(
    doc,
    "소멸 배치",
    "LoyaltyPointExpirationBatch — 배치 크기 200 레코드\n"
    "LoyaltyPointExpirationScheduler — 매일 01시 30분 자동 실행\n"
    "만료 대상일: Date.today().addDays(-1) — 전일 기준으로 만료된 포인트 처리",
    kind="info",
)
add_para(doc, "■ 만료 처리 로직", bold=True, after=2)
add_bullet(doc, "쿼리 조건: EventType='Credit' AND ExpiryDate = expiryDate")
add_bullet(doc, "남은 포인트 = 적립 − Traceability(차감/취소) 합산")
add_bullet(doc, "남은 포인트 ≤ 0 이면 소멸 스킵")
add_bullet(doc, "남은 포인트 > 0 이면 LoyaltyLedger에 EventType='Expiry' 레코드 직접 INSERT")
add_para(doc, "")
add_para(doc, "■ ERP 동기화", bold=True, after=2)
add_bullet(doc, "Operation: POINT_EXPIRY (ErpPointJob)")
add_bullet(doc, "DTO: cd_save_ledger(원 적립 레저), pt_destory(소멸 포인트), cd_use_ledger(소멸 레저), id_insert")

add_heading(doc, "2-7. 마이그레이션 이전 판매 건 처리 — 기타포인트(04)", level=2)
add_callout(
    doc,
    "기타포인트의 정의",
    "'기타포인트(PointTypeCode='04')'는 마이그레이션으로 인한 시스템 오픈 일자 이전에 발생한 판매 건에 대한 포인트 적립에 사용되는 전용 SubType입니다.\n"
    "구 시스템에서 신규 시스템으로 데이터를 이관하면서, 옛 일자 판매 건에 대해 신규 시스템에서 추정 적립을 처리할 때 일반 구매포인트와 분리하기 위해 별도 SubType을 부여한 것입니다.",
    kind="note",
)
add_para(doc, "■ 처리 흐름 (GdPointServiceImpl.creditOptionCPurchase 등)", bold=True, after=2)
add_numbered(doc, "옛 일자 판매 건 입력 시: 등급 보상 중 PT='01'(구매포인트) 비율만 사용 (사은포인트 등 다른 보상은 제외)")
add_numbered(doc, "산출된 적립 포인트의 SubType을 강제로 '기타포인트'로 redirect")
add_numbered(doc, "ERP 전송 시 cd_type_point = '04'로 분류")
add_para(doc, "")
add_para(doc, "■ 차감 시 주의 사항", bold=True, after=2)
add_bullet(doc, "마이그레이션으로 이관된 묶음 잔액은 PointTypeCode='04'(기타포인트)에만 존재")
add_bullet(doc, "안전 정책: 이관 데이터의 부당 차감을 방지하기 위해, 기타포인트 묶음 잔액은 별도 차감 로직(GdOptionCPointAdjustmentUtil)으로 관리")
add_bullet(doc, "신규 판매에서 적립이 누락된 것으로 잘못 추정해 기타포인트 lump에서 강제 차감되는 사례를 방지")
add_para(doc, "")
add_callout(
    doc,
    "인수인계 메모",
    "코드 주석상으로 '옵션 C' 또는 'Option C'라는 용어가 등장하지만, 이는 마이그레이션 정책의 내부 코드명일 뿐입니다. "
    "운영 문맥에서는 '마이그레이션 오픈 일자 이전 판매 건의 포인트 처리'로 이해하시면 됩니다.\n"
    "관련 클래스: GdMigrationUtil, GdOptionCPointAdjustmentUtil, GdPointServiceImpl.creditOptionCPurchase",
    kind="note",
)

add_heading(doc, "2-8. 원장(Ledger) 구조 및 집계 — 인수인계 핵심", level=2)
add_para(
    doc,
    "포인트 운영의 모든 흔적은 아래 5개의 핵심 객체에 기록됩니다. 적립·차감·소멸·잔액 조회 시 "
    "어느 객체를 읽어야 하는지 정확히 알아야 데이터 추적과 디버깅이 가능합니다. 본 절은 포인트 인수인계 시 가장 중요한 부분입니다.",
)

add_callout(
    doc,
    "한눈 요약 — '어떤 작업에 어떤 객체가 쓰이나'",
    "  · 적립(Credit) 시:  LoyaltyLedger(Credit) + LoyaltyMemberCurrency 갱신 + LoyaltyAggrSubPointExprLedger__c 자동 생성/갱신 + TJ.PostPointBalance__c 업데이트\n"
    "  · 차감(Debit) 시:  LoyaltyLedger(Debit) + LoyaltyLedgerTraceability(Debit) 으로 원 적립과 연결 + LoyaltyMemberCurrency.TotalPointsRedeemed 증가 + TJ.PostPointBalance__c 업데이트\n"
    "  · 적립 취소(반품 등):  LoyaltyLedgerTraceability(ActionType='CreditCancellation') 생성\n"
    "  · 차감 취소(사용 취소):  LoyaltyLedgerTraceability(ActionType='DebitCancellation') 생성\n"
    "  · 소멸(Expiry) 시:  LoyaltyLedger(Expiry) + LoyaltyMemberCurrency.TotalPointsExpired 증가\n"
    "  · 회원 현재 총 잔액 조회:  LoyaltyMemberCurrency.PointsBalance (Salesforce 표준 필드, 단일 값)\n"
    "  · 거래 직후 잔액 스냅샷:  TransactionJournal.PostPointBalance__c (거래 단위)\n"
    "  · 만료일별·SubType별 잔액 조회(도움창 등):  LoyaltyAggrSubPointExprLedger__c (커스텀 집계)",
    kind="info",
)

# ----- (1) LoyaltyLedger -----
add_heading(doc, "(1) LoyaltyLedger — 원장의 단일 진실 공급원(Single Source of Truth)", level=3)
add_para(
    doc,
    "모든 포인트 이동은 LoyaltyLedger에 한 줄(Row)로 기록됩니다. EventType 필드로 적립/차감/소멸을 구분하며, "
    "이 객체에 기록되지 않은 포인트 변경은 시스템상 존재하지 않는 것과 같습니다.",
)
add_table(
    doc,
    headers=["필드", "설명"],
    rows=[
        ["Id", "원장 레코드의 고유 ID (Traceability·집계의 키로 사용됨)"],
        ["EventType", "'Credit' = 적립 / 'Debit' = 차감(사용) / 'Expiry' = 소멸"],
        ["Points", "적립·차감·소멸된 포인트 수량 (음수가 아닌 양수로 기록)"],
        ["ExpiryDate", "해당 적립분의 만료 예정일 (소멸 배치가 이 날짜 기준으로 처리)"],
        ["LoyaltyProgramMemberId", "회원(LoyaltyProgramMember) FK"],
        ["LoyaltyProgramCurrencyId", "통화('포인트') FK"],
        ["LoyaltyPgmCrcySubtypeId", "포인트 SubType FK (01 구매 / 02 기념일 / 03 승급 / 04 기타 / 05 에코 / 06 사은 / 07 이벤트)"],
        ["TransactionJournalId", "이 원장이 어떤 거래/저널에서 발생했는지 역추적 FK"],
        ["ActivityDate", "활동 발생일"],
        ["CreatedDate", "시스템 생성 시각"],
    ],
    col_widths=[6, 10],
)
add_callout(
    doc,
    "사용 시점",
    "  · 회원의 모든 포인트 적립·차감·소멸 이력을 시간순으로 조회할 때\n"
    "  · 특정 거래(TransactionJournalId)에서 발생한 포인트 변동을 추적할 때\n"
    "  · 만료 예정 포인트 산출 시 EventType='Credit' AND ExpiryDate <= ? 쿼리로 대상 선별",
    kind="note",
)

# ----- (2) LoyaltyLedgerTraceability -----
add_heading(doc, "(2) LoyaltyLedgerTraceability — 적립-차감 연결 추적", level=3)
add_para(
    doc,
    "차감(또는 취소)이 발생할 때, '어떤 적립분이 얼마나 사용되었는지'를 명시적으로 연결해 주는 객체입니다. "
    "1건의 적립(Credit Ledger)이 여러 번에 나누어 차감될 수 있고, 1건의 차감이 여러 적립분을 동시에 소비할 수 있으므로 M:N 관계를 표현합니다.",
)
add_table(
    doc,
    headers=["필드", "설명"],
    rows=[
        ["CreditLoyaltyLedgerId", "원 적립 LoyaltyLedger ID (Credit 쪽)"],
        ["DebitLoyaltyLedgerId", "차감 LoyaltyLedger ID (Debit 쪽)"],
        ["Points", "이번 매칭에서 차감(또는 취소)된 포인트 수량"],
        ["ActionType", "'Debit' / 'CreditCancellation' / 'DebitCancellation'"],
    ],
    col_widths=[6, 10],
)
add_para(doc, "■ ActionType 값별 의미", bold=True, after=2)
add_bullet(doc, "'Debit': 정상 차감 — 적립분에서 포인트가 사용됨")
add_bullet(doc, "'CreditCancellation': 적립 취소 — 반품 등으로 이미 적립된 포인트를 되돌림 (원 Credit Ledger를 무효화)")
add_bullet(doc, "'DebitCancellation': 차감 취소 — 사용했던 포인트가 다시 회원에게 복귀 (Debit Ledger를 무효화)")
add_callout(
    doc,
    "사용 시점",
    "  · '이 적립분(Credit)에서 얼마나 사용되었나' 잔량 계산: CreditLoyaltyLedgerId 기준 Traceability 합산\n"
    "  · 도움창에서 사용 가능 포인트 표시: Credit Ledger.Points − Σ Traceability.Points(취소 분 제외)\n"
    "  · 반품·취소 흐름 검증: ActionType별 레코드 존재 여부 점검",
    kind="note",
)

# ----- (3) TransactionJournal -----
add_heading(doc, "(3) TransactionJournal — 거래 단위 저널", level=3)
add_para(
    doc,
    "주문/판매/반품/관리자 수동지급 등 '거래 1건'을 표현하는 상위 객체입니다. "
    "각 거래 저널은 1개 이상의 LoyaltyLedger 레코드를 발생시키며, 해당 거래 직후의 잔액 스냅샷을 PostPointBalance__c에 보관합니다.",
)
add_table(
    doc,
    headers=["필드", "설명"],
    rows=[
        ["OrderId / SaleId__c", "원천 주문/판매 FK (역추적용)"],
        ["JournalTypeId", "저널 유형 (Accrual / Redemption / Accrual Reversal 등)"],
        ["JournalSubTypeId", "저널 하위 유형 (Purchase Accrual, Promotion Point Accrual, Manual Coupon Issuance 등)"],
        ["TransactionAmount", "거래 금액"],
        ["PostPointBalance__c", "이 거래 직후의 회원 포인트 잔액 (커스텀 필드, 거래 단위 스냅샷)"],
    ],
    col_widths=[6, 10],
)
add_callout(
    doc,
    "PostPointBalance__c 핵심 메모",
    "이 필드는 거래 처리 로직(GdPointServiceImpl) 내부에서 명시적으로 계산하여 기록합니다.\n"
    "  · 적립 후: tj.PostPointBalance__c = lmcList[0].PointsBalance  (적립 반영 후 LMC 잔액)\n"
    "  · 차감 후: tj.PostPointBalance__c = currentBalance − tj.TransactionAmount\n"
    "  → 거래 단위로 '그 시점의 잔액'을 보고 싶을 때는 이 필드를 보면 됩니다. (현재 잔액은 LMC.PointsBalance)",
    kind="warn",
)

# ----- (4) LoyaltyMemberCurrency -----
add_heading(doc, "(4) LoyaltyMemberCurrency (LMC) — 회원별 누적 통계 & 현재 잔액", level=3)
add_para(
    doc,
    "회원 + 통화('포인트') 조합으로 1개의 레코드가 존재하는 집계 객체입니다. "
    "포인트의 '현재 총 잔액'과 '누적 적립/차감/소멸 총량'을 보유합니다. 회원의 포인트 보유 현황을 가장 빠르게 조회할 수 있는 진입점입니다.",
)
add_table(
    doc,
    headers=["필드", "설명"],
    rows=[
        ["LoyaltyMemberId", "회원(LoyaltyProgramMember) FK"],
        ["LoyaltyProgramCurrencyId", "통화('포인트') FK"],
        ["PointsBalance", "★ 현재 사용 가능한 포인트 총 잔액 (Salesforce 표준 필드, 단일 값)"],
        ["TotalPointsAccrued", "누적 적립 총량"],
        ["TotalPointsRedeemed", "누적 차감(사용) 총량 — 차감 시마다 += tj.TransactionAmount"],
        ["TotalPointsExpired", "누적 소멸 총량"],
    ],
    col_widths=[6, 10],
)
add_callout(
    doc,
    "★ 회원 총 포인트 잔액 조회는 LoyaltyMemberCurrency.PointsBalance ★",
    "관리자 화면에서 '회원 360 > 보유 혜택 > 포인트' 카드에 표시되는 잔액은 이 값을 직접 가져옵니다.\n"
    "LoyaltyLedger를 EventType별로 SUM해서 직접 계산하지 마세요. PointsBalance는 Salesforce 표준 엔진이 자동으로 유지·갱신하는 정합성 보장 필드입니다.",
    kind="warn",
)
add_callout(
    doc,
    "TotalPointsRedeemed 운영 메모",
    "차감 메서드(debitPoint, debitSpecificPoints 등)는 차감 처리 직후 LMC.TotalPointsRedeemed에 거래 금액을 누적 가산합니다.\n"
    "GdPointServiceImpl 코드 예: lmc.TotalPointsRedeemed = (기존값 ?? 0) + tj.TransactionAmount",
    kind="note",
)

# ----- (5) LoyaltyAggrSubPointExprLedger -----
add_heading(doc, "(5) LoyaltyAggrSubPointExprLedger__c — 만료일별·SubType별 집계 원장", level=3)
add_para(
    doc,
    "회원 × 포인트 SubType × 만료일 단위로 적립·차감·소멸 잔량을 미리 집계해 두는 커스텀 객체입니다. "
    "'쿠폰/포인트 적용 도움창'처럼 만료 임박 포인트와 SubType별 잔량을 빠르게 보여줘야 하는 화면에서 이 객체를 우선 조회합니다.",
)
add_table(
    doc,
    headers=["필드", "설명"],
    rows=[
        ["LoyaltyProgramMemberId__c", "회원 FK"],
        ["CurrencySubTypeId__c", "포인트 SubType FK (구매/사은/이벤트/에코/기타/기념일/승급)"],
        ["ExpiryDate__c", "만료일 (같은 만료일·SubType 단위로 묶음)"],
        ["AccruedPoints__c", "해당 묶음 기준 누적 적립 포인트"],
        ["RedeemedPoints__c", "해당 묶음 기준 누적 차감 포인트"],
        ["ExpiredPoints__c", "해당 묶음 기준 누적 소멸 포인트"],
    ],
    col_widths=[6, 10],
)
add_para(doc, "■ 집계 재구축 로직 (GdPointAggregationUtil.rebuildLoyaltyAggrSubPointExprLedger)", bold=True, after=2)
add_bullet(doc, "멤버별 LoyaltyLedger 전체 로드 (LIMIT 50000)")
add_bullet(doc, "EventType별 분류: Credit → AccruedPoints__c / Debit+Traceability → RedeemedPoints__c / Expiry → ExpiredPoints__c")
add_bullet(doc, "저널 SubType별 추가 처리 (Purchase Accrual, Promotion Point Accrual, Point Redeem Cancel 등)")
add_callout(
    doc,
    "사용 시점",
    "  · 만료 임박순 정렬로 도움창에 노출할 포인트 묶음을 조회할 때\n"
    "  · SubType별·만료일별 잔량 보고서를 만들 때\n"
    "  · 정합성 의심 시 GdPointRefreshBatch / GdPointRefreshApplyService 로 LoyaltyLedger를 기준 삼아 재구축",
    kind="note",
)

# ----- 매핑 표 -----
add_heading(doc, "■ 조회 목적별 'Where to Look' — 어디를 봐야 할지 한눈에", level=3)
add_table(
    doc,
    headers=["조회 목적", "참조 객체 · 필드", "비고"],
    rows=[
        ["회원의 현재 포인트 총 잔액", "LoyaltyMemberCurrency.PointsBalance", "표준 필드, 정합성 자동 유지 (★ 1순위)"],
        ["회원의 누적 적립·차감·소멸 총량", "LoyaltyMemberCurrency.TotalPointsAccrued / Redeemed / Expired", "회원 단위 집계"],
        ["거래 직후 잔액 스냅샷", "TransactionJournal.PostPointBalance__c", "거래 단위 시점 잔액 (커스텀)"],
        ["회원의 모든 포인트 이력(시간순)", "LoyaltyLedger WHERE LoyaltyProgramMemberId = ?", "EventType으로 적립/차감/소멸 필터"],
        ["특정 적립분의 남은 사용 가능 잔량", "Credit LoyaltyLedger.Points − Σ Traceability.Points (CreditCancellation 제외)", "Traceability 합산으로 차감액 산출"],
        ["적립-차감 연결 추적", "LoyaltyLedgerTraceability (CreditLoyaltyLedgerId · DebitLoyaltyLedgerId · ActionType)", "ActionType: Debit / CreditCancellation / DebitCancellation"],
        ["만료일·SubType별 잔량 (도움창)", "LoyaltyAggrSubPointExprLedger__c", "사전 집계로 빠른 조회"],
        ["특정 거래에서 발생한 포인트 변동", "LoyaltyLedger WHERE TransactionJournalId = ?", "거래 → 원장 역추적"],
        ["적립 포인트의 SubType (구매/사은/이벤트 등)", "LoyaltyLedger.LoyaltyPgmCrcySubtype.PointTypeCode__c", "01~07 코드"],
    ],
    col_widths=[5.5, 6.5, 4],
)

add_callout(
    doc,
    "디버깅 체크리스트 (포인트 정합성 의심 시)",
    "1) LMC.PointsBalance 와 LoyaltyLedger(Credit − Debit − Expiry) 합산값 비교 → 불일치 시 정합성 깨짐\n"
    "2) Traceability 합산 ≠ Debit Ledger.Points → 적립-차감 매칭 누락\n"
    "3) AggrSubPointExprLedger__c 의 AccruedPoints__c ≠ Credit LoyaltyLedger 합 → GdPointRefreshBatch 로 재집계 필요\n"
    "4) TJ.PostPointBalance__c 와 그 시점 LMC.PointsBalance 가 다르다면 거래 처리 로직 실패 가능성",
    kind="warn",
)

add_heading(doc, "2-9. ERP 연동 흐름", level=2)
add_table(
    doc,
    headers=["연동 유형", "API / 클래스", "Operation 코드"],
    rows=[
        ["적립", "LoyaltyPointService.creditPoints()", "POINT_CREDIT"],
        ["차감", "LoyaltyPointService.debitPoints()", "POINT_DEBIT"],
        ["소멸", "ErpPointJob", "POINT_EXPIRY"],
        ["프로모션 포인트 한도", "ErpPromotionPointJob", "CREATE/UPDATE/DELETE_PROMOTION_POINT"],
        ["Flow 통합", "ErpPointFlowAction", "POINT_CREDIT (Flow)"],
    ],
    col_widths=[4, 7, 5],
)
add_callout(
    doc,
    "ERP 적립 유형 코드 매핑 (ErpPointReqDto)",
    "cd_type_save 적립 유형:  01=구매포인트 / 02=기념일포인트 / 03=승급포인트 / 04=기타포인트 / 05=에코포인트 / 06=사은포인트 / 07=이벤트포인트",
    kind="info",
)

add_page_break(doc)

# ============================================================
# 3. 프로모션 정책
# ============================================================
add_heading(doc, "3. 프로모션 정책", level=1)
add_para(
    doc,
    "프로모션은 특정 기간·매장·상품·회원 등급 등의 조건하에서 할인·쿠폰·포인트 혜택을 제공하는 마케팅 단위입니다. "
    "프로모션 마스터를 중심으로 상품·매장·쿠폰·포인트 자식 엔티티가 연결되며, 모든 변경은 트리거 기반으로 ERP에 자동 동기화됩니다.",
)

add_heading(doc, "3-1. 프로모션 구성 요소", level=2)
add_table(
    doc,
    headers=["엔티티", "설명", "핵심 필드"],
    rows=[
        ["Promotion (마스터)", "프로모션 헤더", "PromotionNo__c, StartDate, EndDate, BrandCode__c, TypeCode__c"],
        ["PromotionProduct", "대상 상품", "ProductId, PromotionId"],
        ["PromotionAccount__c", "대상 매장", "AccountId__c, PromotionId__c, IsActive__c"],
        ["VoucherLimit__c", "프로모션 쿠폰", "VoucherDefinitionId__c, PromotionId__c, Type__c"],
        ["PromotionPointLimit__c", "프로모션 포인트", "PromotionId__c, Value__c, LoyaltyPgmCurrencySubtypeId__c"],
    ],
    col_widths=[4.5, 4, 7.5],
)

add_heading(doc, "3-2. 적용 조건", level=2)
add_para(doc, "■ 기간 조건 — 조회 기준은 '게시 기간(Publish)'", bold=True, after=2)
add_callout(
    doc,
    "두 종류의 기간 필드 — 시작/종료 일자 vs 게시 기간",
    "  · StartDate / EndDate (시작 일자 / 종료 일자): 프로모션 자체의 운영 기간을 의미하는 표준 필드. ERP 응답에서 DT_START / DT_END로 전달됨.\n"
    "  · PublishStartDate__c / PublishEndDate__c (게시 시작 날짜 / 게시 종료 날짜): ★ 프로모션이 고객·매장에 노출(게시)되는 기간. 조회·활성화 판정의 기준 필드.\n\n"
    "프로모션 신규 생성 시에는 두 쌍이 같은 값으로 함께 입력되지만, 운영 중에는 게시 기간만 별도 조정 가능합니다.\n"
    "코드 근거: GdPromotionServiceImpl.cls(생성 라인 49–55, 조회 라인 107–118), PromotionResDTO.cls(활성 판정 라인 75–79)",
    kind="info",
)
add_bullet(doc, "프로모션 활성화 플래그: IsActive = TRUE (현재 운영 필수 조건)")
add_bullet(doc, "활성화 상태(YN_ACTIVE) 판정식: IsActive = TRUE AND PublishStartDate__c ≤ TODAY() ≤ PublishEndDate__c — 두 조건을 모두 만족해야 '활성'")
add_bullet(doc, "조회(검색) 기준: 게시 기간(Publish) 필드를 본다 — StartDate / EndDate가 아님")
add_para(doc, "")
add_para(doc, "■ 기간 Overlap(겹침) 조건 — 정확한 SOQL 표현", bold=True, after=2)
add_callout(
    doc,
    "조회 SOQL (요청에 기간을 넘긴 경우)",
    "  WHERE IsActive = TRUE\n"
    "    AND PublishEndDate__c   >= :요청시작일   ← startDate 파라미터가 들어오면 추가\n"
    "    AND PublishStartDate__c <= :요청종료일   ← endDate 파라미터가 들어오면 추가\n\n"
    "→ 의미: '요청 기간[요청시작, 요청종료]'과 '프로모션 게시 기간[PublishStart, PublishEnd]'이 한 점이라도 겹치면 매칭.\n"
    "→ 요청 기간을 생략하면 Overlap 조건은 추가되지 않으며 IsActive=TRUE 인 모든 프로모션이 조회됩니다.",
    kind="info",
)
add_para(doc, "")
add_para(doc, "■ 매장 조건 (3가지 매칭 방식)", bold=True, after=2)
add_numbered(doc, "전체 매장: IsAllStore__c = TRUE")
add_numbered(doc, "개별 매장: PromotionAccount__c 조회 → StoreCode__c 매칭 ('관련됨 > 적용 가능 매장'에 등록된 개별 매장)")
add_numbered(doc, "매장 그룹(구분/계열): Promotion.StoreType__c(다중선택) ⊇ Account.StoreType__c — '세부 사항 > 매장(구분)/매장(계열)' 드롭다운 선택")
add_para(doc, "")
add_para(doc, "■ 상품 조건", bold=True, after=2)
add_bullet(doc, "전체 상품: IsAllItem__c = TRUE")
add_bullet(doc, "특정 상품: HasTargetProductCondition__c = TRUE + PromotionProduct 조회")
add_para(doc, "")
add_para(doc, "■ 할인율 적용 조건", bold=True, after=2)
add_bullet(doc, "IsDiscountRateApplicable__c — 할인율 적용 가능 여부 제어 플래그")

add_heading(doc, "3-3. 쿠폰 사용 여부 체크 항목", level=2)
add_para(doc, "프로모션의 쿠폰 사용 가능 여부는 다음 4개 플래그의 OR 조건으로 결정됩니다.")
add_callout(
    doc,
    "쿠폰 사용 가능 판정",
    "YN_USE_COUPON = 'Y' ⇐ IsCouponAllowed__c OR IsFixedAmountCouponAllowed__c OR IsRateCouponAllowed__c OR IsSpecificCouponAllowed__c",
    kind="info",
)
add_table(
    doc,
    headers=["필드", "의미", "쿠폰 유형 매핑"],
    rows=[
        ["IsCouponAllowed__c", "전체 쿠폰 허용", "모든 VoucherDefinition Type"],
        ["IsFixedAmountCouponAllowed__c", "정액 쿠폰 허용", "FixedValue"],
        ["IsRateCouponAllowed__c", "정률 쿠폰 허용", "DiscountPercentage"],
        ["IsSpecificCouponAllowed__c", "특정 쿠폰 허용", "VoucherLimit__c(Type__c='PromotionIncludeList')"],
    ],
    col_widths=[6, 4, 6],
)
add_callout(
    doc,
    "핵심 주의",
    "프로모션에 쿠폰을 VoucherLimit__c로 연결해 두었더라도, IsSpecificCouponAllowed__c 체크박스가 체크되지 않으면 해당 쿠폰은 도움창 후보에 오르지 않습니다. "
    "쿠폰 연결과 체크박스 활성화는 반드시 함께 설정되어야 효력이 발생합니다.",
    kind="warn",
)

add_heading(doc, "3-4. 프로모션 포인트 적립·차감 규칙", level=2)
add_bullet(doc, "프로모션 포인트 사용 여부: Promotion.IsPointUsageAllowed__c")
add_bullet(doc, "포인트 한도: PromotionPointLimit__c 자식 레코드로 관리")
add_bullet(doc, "핵심 필드: Type__c(포인트 유형), Value__c(포인트값), LoyaltyPgmCurrencySubtypeId__c(통화 서브타입)")
add_bullet(doc, "응답 매핑: NM_PROM_POINT, NM_TYPE_PROM_POINT, NM_TYPE_POINT, PT_SAVE")
add_bullet(doc, "별첨 매핑: '프로모션 포인트 지급' = 프로모션에 적용된 포인트 적립, '프로모션 포인트 지급 취소' = 적립 취소")

add_heading(doc, "3-5. 적립 불가 포인트 유형 설정", level=2)
add_callout(
    doc,
    "PointTypeCode__c (멀티픽리스트)",
    "프로모션 단위로 '이 프로모션에서는 특정 종류의 포인트는 적립하지 않음'을 지정할 수 있습니다. "
    "PointTypeCode__c 필드는 멀티픽리스트(세미콜론 구분)로, 적립 불가 유형을 선택합니다.",
    kind="info",
)
add_table(
    doc,
    headers=["설정 값", "의미", "ERP 전송 필드"],
    rows=[
        ["'01' 포함", "구매포인트 적립 불가", "yn_pur_point = 'N'"],
        ["'06' 포함", "사은포인트 적립 불가", "yn_gift_point = 'N'"],
        ["'07' 포함", "이벤트포인트 적립 불가", "yn_event_point = 'N'"],
        ["NULL", "모든 포인트 유형 적립 가능", "전부 'Y'"],
    ],
    col_widths=[3, 7, 6],
)
add_para(doc, "관련 클래스: ErpPromotionReqDto.isPointTypeAllowed() — 멀티픽리스트 값에 특정 코드 포함 여부 검사")

add_heading(doc, "3-6. 트리거 및 자동화", level=2)
add_para(doc, "각 자식 엔티티의 변경은 트리거 핸들러를 통해 ERP로 비동기 전파됩니다.")
add_table(
    doc,
    headers=["트리거 핸들러", "처리 이벤트", "ERP Operation"],
    rows=[
        ["PromotionTriggerHandler", "afterInsert / Update / Delete (마스터)", "CREATE/UPDATE/DELETE Master"],
        ["PromotionProductTriggerHandler", "afterInsert / Update / Delete (상품, 200건 청크)", "CREATE/UPDATE/DELETE_PROMOTION_PRODUCT"],
        ["PromotionStoreTriggerHandler", "Insert/Update(=Delete+Create)/Delete", "CREATE/UPDATE/DELETE_PROMOTION_STORE"],
        ["PromotionCouponTriggerHandler", "VoucherLimit__c 변경", "CREATE/UPDATE/DELETE_PROMOTION_COUPON"],
        ["PromotionPointTriggerHandler", "PromotionPointLimit__c 변경", "CREATE/UPDATE/DELETE_PROMOTION_POINT"],
    ],
    col_widths=[6, 5, 5],
)

add_heading(doc, "3-7. ERP 연동 및 업로드", level=2)
add_para(doc, "■ ERP 공통 설정", bold=True, after=2)
add_bullet(doc, "Named Credential: GoldenDewApiNC")
add_bullet(doc, "Timeout: 60초, 최대 재시도 2회")
add_bullet(doc, "토큰 캐시 키: GOLDENDEW (스큐 60초)")
add_para(doc, "")
add_para(doc, "■ 업로드 기능", bold=True, after=2)
add_table(
    doc,
    headers=["업로드 유형", "필수 컬럼", "처리"],
    rows=[
        ["프로모션 매장 업로드 (PromotionAccountUploadControllerV1)", "StoreCode", "CSV 최대 5,000건 권장 → PromotionAccount__c 생성"],
        ["프로모션 상품 업로드 (PromotionProductUploadControllerV1)", "ProductCode", "CSV 최대 5,000건 권장 → PromotionProduct 생성"],
        ["포인트 그룹 지급", "NO_MBR", "CSV 최대 10,000건 권장, 만료 기본 2년"],
        ["쿠폰 그룹 지급", "NO_MBR", "CSV 최대 10,000건 권장"],
        ["쿠폰 한도 정책 (제품)", "ProductCode", "CSV 최대 5,000건 권장"],
        ["쿠폰 한도 정책 (매장)", "StoreCode", "CSV 최대 5,000건 권장"],
    ],
    col_widths=[7.5, 3, 5.5],
)
add_callout(
    doc,
    "업로드 권장 한도",
    "쿠폰 사용 정책의 사용 가능 제품·매장·제품 분류 항목은 각각 최대 5,000건까지 등록 가능합니다. "
    "초과 시 시스템 처리 한도로 인해 조건 설정이 실패하거나 일부만 저장될 수 있습니다.",
    kind="warn",
)

add_page_break(doc)

# ============================================================
# 4. 쿠폰 정책 (PDF 통합)
# ============================================================
add_heading(doc, "4. 쿠폰 정책 — 적용 도움창 조회 기준", level=1)
add_para(
    doc,
    "쿠폰 적용 도움창에는 고객이 보유한 모든 쿠폰이 표시되는 것이 아니라, 현재 거래에서 실제로 사용할 수 있는 쿠폰만 조회됩니다. "
    "고객이 쿠폰을 보유하고 있더라도 아래 조건 중 하나라도 맞지 않으면 도움창에 표시되지 않습니다.",
)

add_heading(doc, "4-1. 기본 조회 조건", level=2)
add_para(doc, "쿠폰이 도움창에 조회되려면 우선 아래 기본 조건을 모두 만족해야 합니다.")
add_bullet(doc, "고객이 실제로 보유한 쿠폰이어야 함")
add_bullet(doc, "아직 사용하지 않은 쿠폰이어야 함")
add_bullet(doc, "유효기간이 지나지 않은 쿠폰이어야 함")
add_bullet(doc, "현재 거래 조건에 맞는 쿠폰이어야 함")
add_bullet(doc, "현재 거래에서 이미 선택한 쿠폰은 제외됨")
add_para(doc, "")
add_callout(
    doc,
    "예시",
    "한 번의 거래에서 같은 쿠폰을 중복으로 사용할 수 없도록, 이미 선택된 쿠폰은 조회 목록에서 제외됩니다.",
    kind="info",
)

add_heading(doc, "4-2. 프로모션이 없는 경우", level=2)
add_para(doc, "현재 거래에 프로모션이 적용되지 않은 경우에는 프로모션 조건을 확인하지 않습니다. 이 경우에는 고객이 보유한 쿠폰 중에서 아래 조건만 확인합니다.")
add_bullet(doc, "쿠폰 사용 가능 상태 여부")
add_bullet(doc, "쿠폰 유효기간")
add_bullet(doc, "상품 조건 / 매장 조건 / 회원 등급 조건 / 거래 구분 조건")
add_bullet(doc, "금액, 수량, 할인율 조건")
add_para(doc, "")
add_para(doc, "즉, 프로모션이 없는 거래에서는 쿠폰 자체에 설정된 사용 조건만 기준으로 조회 여부가 결정됩니다.")

add_heading(doc, "4-3. 프로모션이 있는 경우", level=2)
add_para(doc, "현재 거래에 프로모션이 적용되어 있다면, 먼저 해당 프로모션에서 쿠폰 사용을 허용하는지 확인합니다. 프로모션 설정에 따라 조회되는 쿠폰 범위가 달라집니다.")

add_heading(doc, "(1) 프로모션의 쿠폰 사용 여부 체크 항목", level=3)
add_table(
    doc,
    headers=["체크 항목", "조회 범위"],
    rows=[
        ["전체 쿠폰 사용 여부", "모든 쿠폰이 후보에 오릅니다."],
        ["정액 쿠폰 사용 여부", "정액 쿠폰만 후보에 오릅니다."],
        ["정률 쿠폰 사용 여부", "정률 쿠폰만 후보에 오릅니다."],
        ["특정 쿠폰 사용 여부", "해당 프로모션에 연결된 특정 쿠폰도 후보에 오릅니다."],
    ],
    col_widths=[5, 11],
)

add_heading(doc, "(2) 전체·정액·정률 쿠폰 사용 여부 체크 시", level=3)
add_para(doc, "전체/정액/정률 쿠폰 사용 여부 중 하나라도 체크되어 있고 특정 쿠폰 사용 여부는 체크되지 않은 경우에는, 고객이 보유한 쿠폰 중 해당 유형에 속하는 쿠폰이 후보로 조회됩니다.")
add_callout(
    doc,
    "중요",
    "이 경우에는 프로모션에 직접 연결된 쿠폰만 조회하는 방식이 아닙니다. "
    "특정 쿠폰 사용 여부가 체크된 프로모션이 아니라면, 쿠폰이 특정 프로모션 목록에 연결되어 있지 않더라도 다른 조건이 맞으면 조회될 수 있습니다.",
    kind="info",
)

add_heading(doc, "(3) 특정 쿠폰 사용 여부 체크 시", level=3)
add_para(doc, "프로모션에 특정 쿠폰 사용 여부가 체크되어 있으면, 해당 프로모션에 연결된 쿠폰도 조회 대상에 포함됩니다.")
add_callout(
    doc,
    "핵심 주의",
    "프로모션에 특정 쿠폰을 연결해 두었더라도, 특정 쿠폰 사용 여부 체크박스가 체크되어 있지 않으면 해당 쿠폰은 후보에 오르지 않습니다. "
    "프로모션에 쿠폰을 설정한 것만으로는 효력이 없으며, 반드시 특정 쿠폰 사용 여부가 함께 체크되어 있어야 해당 쿠폰들이 조회 대상에 포함됩니다.",
    kind="warn",
)

add_para(doc, "예시: 프로모션 A에 쿠폰 1과 쿠폰 2가 연결되어 있을 때, 특정 쿠폰 사용 여부 체크 상태에 따라 결과가 달라집니다.")
add_table(
    doc,
    headers=["구분", "특정 쿠폰 사용 여부", "프로모션 A 연결 여부", "조회 결과"],
    rows=[
        ["케이스 1", "체크됨", "쿠폰 1 (연결됨)", "조회됨"],
        ["케이스 2", "체크됨", "쿠폰 2 (연결 안됨)", "조회되지 않음"],
        ["케이스 3", "체크되지 않음", "쿠폰 1 (연결됨)", "후보에 오르지 않음"],
    ],
    col_widths=[3, 4, 5, 4],
)
add_para(doc, "→ 케이스 3에서 보듯이, 쿠폰을 프로모션에 연결해 두는 것만으로는 충분하지 않습니다. 특정 쿠폰 사용 여부가 체크되어야 연결 쿠폰 조건이 효력을 발휘합니다. 이 조건은 일반 쿠폰 허용 조건과 함께 판단됩니다.")

add_heading(doc, "4-4. 쿠폰 사용 조건별 확인 기준", level=2)
add_para(doc, "쿠폰에는 프로모션 외에도 여러 사용 조건이 설정될 수 있습니다. 도움창에서는 현재 거래 정보가 쿠폰 조건에 맞는지 확인한 후 조회 여부를 결정합니다.")

add_para(doc, "■ 상품 조건", bold=True, after=2)
add_para(doc, "쿠폰이 특정 상품 또는 특정 상품군에만 사용 가능하도록 설정되어 있다면, 현재 거래 상품이 해당 조건에 맞아야 조회됩니다.")

add_para(doc, "■ 매장 조건", bold=True, after=2)
add_para(doc, "쿠폰이 특정 매장 또는 특정 매장 유형에서만 사용 가능하도록 설정되어 있다면, 현재 거래 매장이 해당 조건에 맞아야 합니다. 특정 매장과 매장 유형이 모두 설정되어 있는 경우에는 둘 중 하나만 맞아도 사용 가능한 것으로 판단합니다.")
add_table(
    doc,
    headers=["쿠폰 설정값", "내용"],
    rows=[
        ["특정 매장", "20001"],
        ["매장 유형", "백화점"],
    ],
    col_widths=[5, 11],
)
add_para(doc, "→ 현재 매장이 20001이거나 백화점 유형이면 쿠폰이 조회될 수 있습니다.")

add_para(doc, "■ 회원 등급 조건", bold=True, after=2)
add_para(doc, "쿠폰이 특정 회원 등급(GIP / VIP / VVIP / SVIP) 전용으로 설정되어 있다면, 고객의 회원 등급이 해당 조건에 포함되어야 합니다. 예를 들어 VVIP 전용 쿠폰은 일반(GIP) 회원에게 조회되지 않습니다.")

add_para(doc, "■ 거래 구분 조건", bold=True, after=2)
add_para(doc, "쿠폰이 특정 거래 구분에서만 사용 가능하도록 설정되어 있다면, 현재 거래 구분이 해당 조건과 일치해야 합니다.")
add_bullet(doc, "일반 판매에서만 사용 가능")
add_bullet(doc, "수선 거래에서만 사용 가능")
add_bullet(doc, "특정 주문 유형에서만 사용 가능")
add_para(doc, "")
add_callout(doc, "참고", "현재 시스템에서는 거래 구분 값이 요청에 포함된 경우에만 해당 조건을 확인합니다.", kind="info")

add_para(doc, "■ 금액 조건", bold=True, after=2)
add_para(doc, "쿠폰에 최소 구매 금액 또는 금액 제한 조건이 있으면, 현재 거래 금액이 해당 조건을 만족해야 합니다.")
add_bullet(doc, "10만 원 이상 구매 시 사용 가능")
add_bullet(doc, "50만 원 이하 거래에서만 사용 가능")
add_para(doc, "")

add_para(doc, "■ 수량 조건", bold=True, after=2)
add_para(doc, "쿠폰에 구매 수량 조건이 있으면, 현재 거래의 총 상품 수량이 조건을 만족해야 합니다.")
add_bullet(doc, "2개 이상 구매 시 사용 가능")
add_bullet(doc, "1개 구매 시에만 사용 가능")
add_para(doc, "")

add_para(doc, "■ 할인율 조건", bold=True, after=2)
add_para(doc, "쿠폰에 할인율 조건이 있으면, 현재 거래의 할인율이 해당 조건을 만족해야 합니다.")
add_bullet(doc, "할인율 10% 이하일 때만 사용 가능")
add_bullet(doc, "이미 높은 할인이 적용된 거래에는 사용 불가")
add_para(doc, "")

add_heading(doc, "4-5. 쿠폰이 조회되지 않을 때 확인 순서", level=2)
add_numbered(doc, "고객이 해당 쿠폰을 실제로 보유하고 있는지")
add_numbered(doc, "쿠폰이 아직 사용 가능한 상태인지")
add_numbered(doc, "쿠폰 유효기간이 지나지 않았는지")
add_numbered(doc, "현재 거래에서 이미 선택한 쿠폰은 아닌지")
add_numbered(doc, "프로모션에 쿠폰 사용 여부(전체/정액/정률/특정)가 체크되어 있는지")
add_numbered(doc, "프로모션에 특정 쿠폰 사용 여부가 체크된 경우, 해당 쿠폰이 프로모션에 연결되어 있거나 일반 쿠폰 허용 조건에 맞는지")
add_numbered(doc, "프로모션에 쿠폰만 연결해 두고 특정 쿠폰 사용 여부 체크를 빠뜨리지 않았는지")
add_numbered(doc, "현재 거래 상품이 쿠폰의 상품 조건에 맞는지")
add_numbered(doc, "현재 거래 매장이 쿠폰의 매장 조건에 맞는지")
add_numbered(doc, "고객의 회원 등급이 쿠폰 조건에 맞는지")
add_numbered(doc, "현재 거래 구분이 쿠폰 조건에 맞는지")
add_numbered(doc, "거래 금액, 수량, 할인율 조건을 만족하는지")

add_heading(doc, "4-6. 대표 예시", level=2)

add_para(doc, "■ 예시 1. 정액 쿠폰 사용 여부 체크 프로모션", bold=True, after=2)
add_para(doc, "고객이 3만 원 정액 쿠폰을 보유하고 있고, 쿠폰이 사용 전이며 유효기간도 남아 있습니다. 프로모션에서는 정액 쿠폰 사용 여부가 체크되어 있고, 특정 쿠폰 사용 여부는 체크되어 있지 않습니다.")
add_callout(doc, "결과: 해당 쿠폰은 조회될 수 있습니다.", "다만 상품, 매장, 회원 등급, 금액 등 다른 조건도 함께 만족해야 합니다.", kind="ok")

add_para(doc, "■ 예시 2. 특정 쿠폰 사용 여부 체크 프로모션", bold=True, after=2)
add_para(doc, "고객이 3만 원 정액 쿠폰을 보유하고 있더라도, 프로모션에 특정 쿠폰 사용 여부만 체크되어 있고 정액 쿠폰 사용 여부는 체크되어 있지 않으며 해당 쿠폰이 그 프로모션에 연결되어 있지 않은 경우입니다.")
add_callout(doc, "결과: 쿠폰은 조회되지 않습니다.", "특정 쿠폰 사용 여부만 체크된 프로모션에서는 연결되지 않은 쿠폰은 후보에서 제외됩니다.", kind="warn")

add_para(doc, "■ 예시 3. 쿠폰만 연결하고 체크는 빠뜨린 프로모션", bold=True, after=2)
add_para(doc, "프로모션 A에 쿠폰 1이 연결되어 있지만, 특정 쿠폰 사용 여부 체크박스는 체크되어 있지 않은 경우입니다. 고객은 쿠폰 1을 보유하고 있고 사용 전이며 유효기간도 남아 있습니다.")
add_callout(
    doc,
    "결과: 쿠폰 1은 해당 프로모션의 후보에 오르지 않습니다.",
    "프로모션에 쿠폰을 연결해 두었더라도 특정 쿠폰 사용 여부 체크박스가 체크되지 않으면 그 연결 설정은 효력이 없습니다. "
    "이 경우 다른 체크 항목(전체/정액/정률)의 기준에 따라 조회 범위가 결정됩니다.",
    kind="warn",
)

add_para(doc, "■ 예시 4. 매장 조건이 맞지 않는 경우", bold=True, after=2)
add_para(doc, "고객이 쿠폰을 보유하고 있고 쿠폰 상태와 유효기간도 문제가 없지만, 해당 쿠폰이 특정 매장에서만 사용 가능하도록 설정되어 있고 현재 거래 매장이 해당 매장이 아닌 경우입니다.")
add_callout(doc, "결과: 쿠폰은 조회되지 않습니다.", "매장 조건이 일치하지 않는 쿠폰은 도움창에 노출되지 않습니다.", kind="warn")

add_page_break(doc)

# ============================================================
# 5. 한 줄 요약
# ============================================================
add_heading(doc, "5. 한 줄 요약", level=1)
add_callout(
    doc,
    "통합 요약",
    "골든듀 멤버십·포인트·프로모션·쿠폰 정책은 모두 'Golden Dew' 로열티 프로그램을 중심으로 운영되며, "
    "회원의 3년 누적 구매 실적에 따라 GIP / VIP / VVIP / SVIP 4단계 등급이 매년 2월 1일에 자동 전환되고, "
    "등급별로 구매포인트와 사은포인트가 차등 적립(예: GIP 1%+1%, VIP 1.5%+2%)되며, "
    "에코·이벤트·기념일·승급·기타 포인트 등 총 7종의 SubType이 종류별 우선순위에 따라 사용·소멸됩니다.\n\n"
    "프로모션은 기간·매장·상품·할인 조건에 따라 쿠폰·포인트 혜택을 제공하며, "
    "프로모션 단위로 적립 불가 포인트 유형(구매/사은/이벤트)을 별도 지정할 수 있습니다.\n\n"
    "쿠폰 적용 도움창은 고객이 보유한 쿠폰 중 현재 거래·상품·매장·회원 등급·금액·수량·할인율·프로모션 정책에 모두 맞는 쿠폰만 조회합니다. "
    "특히 프로모션에 특정 쿠폰 사용 여부가 체크된 경우에만, 해당 프로모션에 직접 연결된 쿠폰도 조회 후보에 포함됩니다.",
    kind="info",
)

# 저장
output_path = r"C:\Users\milvus-0\Downloads\골든듀_통합정책본.docx"
doc.save(output_path)
print(f"OK: {output_path}")
